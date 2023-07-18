"""Модуль, в котором я собрал разные функции, классы и прочие приблуды, которыми я пользуюсь в своих программах. Теперь можно его расширять, дополнять, улучшать, модернизировать вообще кому угодно - пишите сюда вообще всё, что хотите и считаете нужным

Итак, что тут есть?

Функции:
    checking_arguments - Проверка типа и значения переменной. Ещё может вывести длину переменной, если подобное возможо для данного типа переменой.

    file_names_csv_and_xlsx - Создание двух переменных-названий файлов с расширениями csv и xlsx.

    from_xlsx_to_csv - Конвертация xlsx в csv.

    from_csv_to_xlsx - Конвертация csv в xlsx .

    line_by_line - Запись строки в csv-файл в его конец.

    decision - Функция, которая возвращает значение в зависимости от вопроса.

    the_new_order - Позаимствованный из интернета код, который двигает туда-сюда листы в экселе.

    csv_cleaning - Функция очистки csv-файла - в целом, морально устарела.

    ideal_message - Печатает строку состояния работы цикла в виде "Сделано столько-то чего-то там, осталось столько-то чего-то там. Процентов: проценты. Прошло времени - столько-то секунд. Осталось столько-то секунд".

    dadata_inn_and_address - Функция, обращающаяся к Дадате и забирающая оттуда всю необходимую информацию по данному ИНН.

    dadata_left - Возвращает количество оставшихся запросов.

    current_database - Возвращает параметры базы данных в PostgreSQL.

    veza_design - Простенький метод, который задаёт созданное Майей оформление для программок.

    temporary_filename - Возвращает имя временного файла, который будет храниться в той же директории, что и программа.

    convert_using_win32 - Конвертация из различных форматов посредством модуля win32 (где конвертация - сохранение файла в нужном формате через заданное приложение)

    ordered_content_from_docx - Функция читает подряд docx-файл и возвращает все основные объекты файла в порядке их появления.

    flatten_dictionary - Функция, которая "выравнивает" словарь, делая его одноуровневым

    _onKeyRelease - функция, которая позволяет копировать, вырезать и вставлять текст в оконной форме даже с кириллической раскладкой

    xlsx_file_beautifulication - делает экселевский файл "красивым" - пытается подогнать ширину столбцов под текст и делает везде все нужные границы

    from_base10_to_baseXX - Перевод числа из десятеричной системы счисления в заданную.

    from_baseXX_to_base10 - Перевод числа из заданной системы счисления в десятеричную.

    find_all_systems - Определяет количество систем, записанных внутри одного обозначения.

    do_something_fun - создаёт анимированный градиентный экран

Классы:
    Blank - Общий класс для данных, которые извлекаются из бланка-заказа.

    DocxExpand - Дополнительные методы или функции, которые должны расширить функционал модуля docx.

    BaseXX - Класс, представляющий собой число в заданной системе счисления.

    Base12 - Класс чисел в 12-ичной системе счисления.
"""

import pandas as pd
import time
import os
from dadata import Dadata
# import PySimpleGUI as sg
import docx
from docx import Document
import re
from docx.text.paragraph import Paragraph
import itertools as it
from io import BytesIO
from docx2txt import process
from tabulate import tabulate

TOKEN = "f9607f7223c70da82d49c81434d14fa7b9ab635e"
SECRET = "c8f87dab214c110e816b26801647a3912ef61762"
ALL_DIGITS = '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZБГДЁЖИЙЛПФЦЧШЩЪЫЬЭЮЯÞΣΨΩ'
SUPPORTED_EXCTENTIONS_FOR_BLANK = ('.docx', '.doc', '.rtf', '.pdf')

# popup_true_false = lambda question:sg.popup_yes_no(question) == 'Yes'

# ===============================================================================================================================================================================================

class PrintDebugModeOn:
    """Своеобразный класс-функция, задача которой выводить на экран информацию тогда и только тогда, когда того требуется. Ведёт себя абсолютно точно как обычный Принт (закидывайте в него аргументы, разделитье и знак окончания строки), но чтобы его включить, надо вызвать метод debug_mode_tumbler(). Ничего не возвращает, не создаёт объект, после принта сразу удаляет самого себя в момент инициализации, не пригоден для того, чтобы это было полноценным объектом, лол
    """
    debug_mode_on=True

    def __init__(self, *args, sep=' ', end='\n') -> None:
        self.args = args
        self.sep = sep
        self.end = end
        if self.debug_mode_on:
            print(self.sep.join(str(arg) for arg in self.args) + self.end)
        del self
        pass

    @classmethod
    def debug_mode_tumbler(self):
        self.debug_mode_on = not self.debug_mode_on

    @classmethod
    def check_debug_mode(self):
        print(self.debug_mode_on)

    def __str__(self) -> str:
        return self.sep.join(self.args) + self.end
    
    def __repr__(self) -> str:
        return 'PrintDebugModeOn(' + ', '.join(f"'{arg}'" if isinstance(arg, str) else str(arg) for arg in self.args) + f", sep={self.sep}, end={self.end})"
        pass
    pass

def checking_arguments(*args):
    """Проверка типа и значения переменной. Ещё может вывести длину переменной, если подобное возможо для данного типа переменой.
    """

    for arg in args:
        print(type(arg), arg, sep=': ')
        try:
            print(len(arg))
        except:
            print('Данная переменная не имеет длинны!')

def file_names_csv_and_xlsx(file:str):
    """Создание двух переменных-названий файлов с расширениями csv и xlsx

    Args:
        file (str): Название файла. Лучше всего писать его без расширения.

    Returns:
        file_csv, file_xlsx: Кортеж, где первый элемент - название файла с расширением csv, а второй - xlsx
    """

    file_name, _ = os.path.splitext(file)
    return file_name + '.csv', file_name + '.xlsx'

def from_xlsx_to_csv(filename_xlsx:str, filename_csv:str, to_print=True):
    """Конвертация xlsx в csv

    Args:
        filename_xlsx (str): xlsx-файл
        filename_csv (str): csv-файл
        to_print (bool, optional): Если нужно вывести сообщение о том, что был записан новый файл, написать True. По умолчанию тут стоит False.
    """

    if to_print:
        print('Начинается конвертация файла', filename_xlsx)
    pd.read_excel(filename_xlsx).to_csv(filename_csv, index=False, header=True)
    if to_print:
        print('Конвертирован файл', filename_xlsx)

def from_csv_to_xlsx(filename_xlsx:str, filename_csv:str, to_print=True):
    """Конвертация csv в xlsx 

    Args:
        filename_xlsx (str): xlsx-файл
        filename_csv (str): csv-файл
        to_print (bool, optional): Если нужно вывести сообщение о том, что был записан новый файл, написать True. По умолчанию тут стоит False.
    """
    
    if to_print:
        print('Начинается конвертация файла', filename_csv)
    pd.read_csv(filename_csv).to_excel(filename_xlsx, index=False, header=True)
    if to_print:
        print('Конвертирован файл', filename_csv)

def line_by_line(string:list, file:str, to_print=False):
    """Запись строки в csv-файл в его конец, переменная messaging должна отвечать за то, нужно ли выводить сообщение о записи

    Args:
        string (list): Строка в виде списка, которую надо записать
        file (str): Название файла
        to_print (bool, optional): Если нужно вывести сообщение о том, что было записано и в какой файл, написать True. По умолчанию тут стоит False.
    """

    pd.DataFrame([string]).to_csv(file, mode='a', header=False, index=False)
    if to_print:
        print(f'В файл {file} было записано: {string}')

def decision(question:str, the_type:type, limit_do=0, limit_up=0, use_pysimplegui=False) -> bool|int|float:
    """Функция, которая возвращает значение в зависимости от вопроса. Может возвращать данные разных типов (лучше логическое или число, типа, для строки можно просто сделать инпут), завичит от того, что ввести. Может осуществляться либо через всплывающие окна, либо через консоль. Если нужно ввести число, то ОБЯЗАТЕЛЬНО надо не забыть про ограничения числа, что ли

    Args:
        question (str): Вопрос, на который надо дать ответ
        the_type (type): Тип переменной, которой надо вернуть
        limit_do (int, optional): Ограничение числового значения снизу. Defaults to 0.
        limit_up (int, optional): Ограничение числового значения сверху. Defaults to 0.
        use_pysimplegui (bool, optional): Нужно ли использовать оконную форму (с ней удобней, правда!). Defaults to False.

    Returns:
        the_type: Нужная переменная нужного типа
    """

    if the_type == bool:
        if use_pysimplegui:
            answer = popup_true_false(question)
        else:
            answer = ''
            while type(answer) != bool:
                answer = input(question + '\n')
                if answer.upper() in ('ДА', 'YES', 'LF', 'НУЫ', 'Y', 'Д', 'TRUE'):
                    answer = True
                elif answer.upper() in ('НЕТ', 'NO', 'YTN', 'ТЩ', 'N', 'Н', 'FALSE'):
                    answer = False
                else:
                    print('Введите ещё раз.')
    else:
        answer = ''
        while not isinstance(answer, the_type):
            # answer = sg.popup_get_text(question) if use_pysimplegui else input(question + '\n')
            if answer is None:
                return None
            try:
                answer = the_type(answer)
            except ValueError:
                err_mess = 'Некорректно введённое значение.'
                # sg.popup(err_mess) if use_pysimplegui else print(err_mess)
            else:
                if not (limit_do <= answer <= limit_up):
                    err_mess = 'Значение за пределами диапазона.'
                    answer = ''
                    # sg.popup(err_mess) if use_pysimplegui else print(err_mess)
    return answer

def the_new_order(file, fpos, tpos):
    from openpyxl import load_workbook
    """Позаимствованный из интернета код, вот его описание: Takes a list of ints, and inserts the fpos (from position) int, to tpos (to position)
    
    Важное примечание: если нужно просто вытащить данные через с разных листов Экселя, пользуйтесь пандасом, у него для этого есть все инструменты!

    Args:
        file (_type_): Сюда, по идее, записывается название файла
        fpos (_type_): Сюда - исходная позиция, откуда двигается
        tpos (_type_): Сюда - то, куда двигать (кажется)
    """

    print('Начинается работа по перемещению листов в файле', file)
    wb = load_workbook(filename=file, data_only=True)
    shlist = wb.sheetnames  # get current order sheets in workbook
    lst = []
    lpos = (len(shlist) - 1) # last position
    if lpos >= fpos > tpos >= 0:  # move from a high to low position
        for x in range(lpos+1):
            if x == tpos:
                lst.append(fpos)
            elif tpos < x <= fpos:
                lst.append(x-1)
            else:
                lst.append(x)
    if lpos >= tpos > fpos >= 0:  # move from a low to high position
        for x in range(lpos+1):
            if x == tpos:
                lst.append(fpos)
            elif fpos <= x < tpos:
                lst.append(x+1)
            else:
                lst.append(x)
    wb._sheets = [wb._sheets[i] for i in lst]  # get each object instance from  wb._sheets, and replace
    wb.save(filename=file)
    print('Заканчивается работа по перемещению листов в файле', file)

def csv_cleaning(file_csv:str):
    """Функция очистки csv-файла - в целом, морально устарела

    Args:
        file_csv (str): название csv-файла
    """

    dataframe=open(file_csv, 'w+')
    dataframe.seek(0)
    dataframe.close

def ideal_message(tek_i:int, all_i:int, edinitsa_mn_chislo_rod_padezh:str, time_start:float, show_left_time=True, return_message=False):
    """Печатает строку состояния работы цикла в виде "Сделано столько-то чего-то там, осталось столько-то чего-то там. Процентов: проценты. Прошло времени - столько-то секунд. Осталось столько-то секунд". Желательно писать это в самом конце рабочего поля цикла.

    Args:
        tek_i (int): текущий номер итерации (уж придумайте, как его задать)
        all_i (int): общее количество ожидаемых итераций (обычно можно просто нахерачить длину массива, но если есть массив в массиве, то включается высшая матемаика)
        edinitsa_mn_chislo_rod_padezh (str): единица измерения в родительном падеже и множественном числе.
        time_start (float): время начала работы цикла. Задать ПЕРЕД всем циклом в формате "переменная = time.time()"
        show_left_time (bool, optional): На случай, если необходимо получить количество оставшегося времени. Эффективно применять, если массив большой, а время работы предполагается гигантским, на малых сроках он работает не особо хорошо. Defaults to True.
        return_message (bool, optional): Надо ли возвращать сообщение как строку. Defaults to False.

    Returns:
        _type_: _description_
    """
    
    all_time = round((time.perf_counter() - time_start), 2)
    procents = round((100 * tek_i / all_i), len(str(all_i)) - 2)

    ideal_message = f'Сделано {tek_i} {edinitsa_mn_chislo_rod_padezh}, осталось {(all_i - tek_i)} {edinitsa_mn_chislo_rod_padezh}. Процентов: {procents}. Прошло времени: {all_time} секунд.'
    if show_left_time:
        ideal_message += f' Осталось {round(all_time / (procents / 100) - all_time, 2)} секунд.'
    return ideal_message if return_message else print(ideal_message)

def dadata_inn_and_address(inn_or_address:str, only_main=True, is_it_inn=True, time_left=30):
    """Функция, обращающаяся к Дадате и забирающая оттуда всю необходимую информацию по данному ИНН

    Args:
        inn_or_address (str): Сюда нужно написать либо ИНН компании, либо адрес. ВАЖНО! Тип - str!
        only_main (bool, optional): Прописать False, если нужно забрать ещё и филиалы. Defaults to True.
        is_it_inn (bool, optional): Прописать False, если нужно обратиться по адресу. Defaults to True.
        time_left (int, optional): Время задержки в случае вылета. Defaults to 30.

    Returns:
        result[0]['data']: словарь всех данных
        result: абсолютно все результаты
        None: если ничего не найдено
    """

    dadata = Dadata(TOKEN)
    yes_result = False
    while yes_result == False: 
        try:
            result = dadata.find_by_id(name="party", query=inn_or_address) if is_it_inn else dadata.suggest("address", inn_or_address)
        except Exception as error:
            print(error)
            while time_left > 0:
                print('Осталось', time_left, 'сек')  
                time.sleep(1)
                time_left -= 1
            print('Новая попытка')
        else:
            yes_result = True
    
    if result != []:
        if only_main:
            return result[0]['data']
        else:
            return result
    else:
        return None

def dadata_left(to_print=False):
    """Возвращает количество оставшихся запросов

    Args:
        to_print (bool, optional): Надо ли сразу печатать оставшееся количество. Defaults to False.

    Returns:
        int: left - количество оставшихся запросов
    """

    with Dadata(TOKEN, SECRET) as dadata:
        left = 100_000 - int(dadata.get_daily_stats()['services']['suggestions'])
        if to_print:
            print('Осталось запросов:', left)
        return left

def current_database():
    """Возвращает параметры базы данных в PostgreSQL

    Returns:
        connection: connect - Параметры базы данных, необходимо присвоить переменой, чтобы потом пользоваться этим соединением
    """
    
    import psycopg2
    return psycopg2.connect(
        user="postgres",
        password="p@ssw0rd",  # пароль, который указали при установке PostgreSQL
        host="192.168.30.223",
        port="5432",
        database="VEZA")

# def veza_design():
#     """Простенький метод, который задаёт созданное Майей оформление для программок
#     """

#     sg.LOOK_AND_FEEL_TABLE['MyCreatedTheme'] = {
#     'BACKGROUND': '#BECBBA',
#     'TEXT': '#172412',
#     'INPUT': '#ECF4E7',
#     'TEXT_INPUT': '#172412',
#     'SCROLL': '#172412',
#     'BUTTON': ('#172412', '#EEFFFF'),
#     'PROGRESS': ('#172412', '#EEFFFF'),
#     'BORDER': 3, 
#     'SLIDER_DEPTH': 2, 
#     'PROGRESS_DEPTH': 2, }  
#     # Switch to use your newly created theme
#     sg.theme('MyCreatedTheme')

def temporary_filename(filename:str, postfix='', new_extension='') -> str:
    """Возвращает имя временного файла, который будет храниться в той же директории, что и программа.

    Args:
        filename (str): имя файла
        postfix (str): постфикс, который будет добавляться к имени файла и содержать в себе в первую очередь расширение. По умолчанию равен '' - в таком случае будет просто создаваться комия файла в папке.

    Returns:
        str: имя временного файла с заданным постфиксом
    """

    if '/' in filename:
        filename = filename.replace('/', '\\')
    filename_short, extention = os.path.splitext(filename)
    filename_short = filename_short.split('\\')[-1]
    postfix = 'копия' if postfix == '' else postfix
    new_extension = extention if new_extension == '' else new_extension
    return f"{os.getcwd()}\\{filename_short} {postfix}{new_extension}"

def convert_using_win32(old_filename:str, new_filename:str, extension:str, to_print=True):
    """Конвертация всего, что можно открыть в Ворде, в docx. Иногда оно можно не сработать по непонятным мне причинам - я попытался это предусмотреть, однако есть ещё вероятность, что что-то может пойти наперекосяк. Модуль win32 - мощный инструмент. Возможно, даже слишком мощный. Не гарантируется работоспособность на других операционных системах

    Args:
        old_filename (str): Имя старого файла
        new_filename (str): Имя нового файла
        extension (str): в какое расширение нужно сконвертировать. Поддерживаемые варианты: 'docx', 'xlsx'
    """
    from win32com import client as wc
    old_filename, new_filename = os.path.realpath(old_filename), os.path.realpath(new_filename)
    if to_print:
        print('Начинается конвертация файла', old_filename)
    match extension:
        case 'docx':
            app = 'Word.Application'
            save_code = 16
        case 'xlsx':
            app = 'Excel.Application'
            save_code = 51
    doc = wc.Dispatch(app).Documents.Open(old_filename)
    if doc is not None:
        doc.SaveAs(new_filename, save_code)
        doc.Close(False)
        if to_print:
            print('Конвертирован файл', old_filename)
    else:
        print('Не удалось сконвертировать файл! Попробуйте сделать вручную!')

def ordered_content_from_docx(value:str|docx.document.Document, save_as_objects=False) -> tuple:
    """Функция читает подряд docx-файл и возвращает все основные объекты файла в порядке их появления. Это серьёзно доработанный код из интернета, из которого удалено всё лишнее и, на мой скромный вкус, не значащее. Исходный код лежит в соседнем файле. Главное улучшение - уход от датафреймовой структуры в пользу более примитивной, но более надёжной структуры двухмерных списков-массивов, что, на мой скромный вкус, лучше. Позволяет, как минимум, уйти от необходимости импортировать функцию по игнорированию ФьючерВорнинг.

    Args:
        value (str|docx.document.Document): вордовский файл или объект
        save_as_objects (bool): надо ли сохранять параметры в список как объекты модуля докс или как текст или ещё что? По умолчанию False

    Raises:
        ValueError: _description_

    Returns:
        tuple: document, combined_list, image_list: объект класса Document для дальнейшей с ним работы и два списка - общий и картиночек

    Yields:
        _type_: каждый параграф и таблицу в порядке появления в документе
    """

    from docx.text.paragraph import Paragraph
    import xml.etree.ElementTree as ET
    from docx.document import Document as doctwo
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    from xml.etree import ElementTree
    from io import StringIO
    import base64

    if isinstance(value, str):
        document = Document(value)
    elif isinstance(value, docx.document.Document):
        document = value
    else:
        raise TypeError("Серьёзно, какого хера вы это ввели, обмудки?")

    ##This function extracts the tables and paragraphs from the document object
    def iter_block_items(parent):
        """Yield each paragraph and table child within *parent*, in document order. Each returned value is an instance of either Table or Paragraph. *parent* would most commonly be a reference to a main Document object, but also works for a _Cell object, which itself can contain paragraphs and tables.

        Args:
            parent (_type_): _description_

        Raises:
            ValueError: _description_

        Yields:
            Paragraph | Table: параграф-объект или таблица-объект
        """
        if isinstance(parent, doctwo):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    combined_list = []
    image_list = []
    i = 0
    imagecounter = 0

    for block in iter_block_items(document):
        if 'text' in str(block):
            isappend = False
            
            runboldtext = ''
            for run in block.runs:                        
                if run.bold:
                    runboldtext = runboldtext + run.text
                    
            style = str(block.style.name)

            if save_as_objects:
                appendtxt = block
            else:
                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n","")
                appendtxt = appendtxt.replace("\r","")
            tabid = 'Novalue'
            
            isappend = True
            for run in block.runs:
                xmlstr = str(run.element.xml)
                my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                if 'pic:pic' in xmlstr:
                    for pic in ET.fromstring(xmlstr) .findall('.//pic:pic', my_namespaces):
                        name_attr = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces).get("name")
                        embed_attr = pic.find("pic:blipFill/a:blip", my_namespaces).get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        isappend = True
                        appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                        image_list.append([
                            imagecounter,
                            embed_attr,
                            name_attr,
                            base64.b64encode(document.part.related_parts[embed_attr]._blob).decode()])
                        style = 'Novalue'
                    imagecounter = imagecounter + 1
                
        elif 'table' in str(block):
            isappend = True
            style = 'Novalue'
            appendtxt = block
            tabid = i
            i += 1
        if isappend:
            combined_list.append([appendtxt, tabid, style])

    return document, combined_list, image_list

def flatten_dictionary(dictionary:dict, use_new_key:bool, parent_key='', sep='_') -> dict:
    """Функция, которая "выравнивает" словарь, делая его одноуровневым. Возможно либо сохранение внутренних названий ключей, либо же полное наименование ключей. Старый словарь не изменяется, создаётся новый словарь

    Args:
        dictionary (dict): словарь, который надо "выравнять"
        use_new_key (bool): надо ли использовать новые ключи
        parent_key (str, optional): "родительский" ключ. Defaults to ''.
        sep (str, optional): Разделитель. Defaults to '_'.

    Returns:
        dict: выровненный словарь
    """
    items = []
    for key, value in dictionary.items():
        new_key = parent_key + sep + key if (parent_key and use_new_key) else key
        items.extend(flatten_dictionary(value, use_new_key, new_key, sep=sep).items()) if isinstance(value, dict) else items.append((new_key, value))
    return dict(items)

def _onKeyRelease(event):
    """Взятый из интернета код, который позволяет копировать, вырезать и вставлять текст в оконной форме даже с кириллической раскладкой

    Args:
        event (_type_): _description_
    """
    ctrl  = (event.state & 0x4) != 0
    if event.keycode==88 and ctrl and event.keysym.lower() != "x":
        event.widget.event_generate("<<Cut>>")
    if event.keycode==86 and ctrl and event.keysym.lower() != "v":
        event.widget.event_generate("<<Paste>>")
    if event.keycode==67 and ctrl and event.keysym.lower() != "c":
        event.widget.event_generate("<<Copy>>")

def progress_bar_updater(progress_bar, c:int, negative=False):
    """Просто обновление счётчика. Увеличивает (и ли уменьшает) его его на один и обновляет прогресс-бар. Вынесено в модуль, потому что слишком часто приходится проделывать эту процедуру в моих программах...
        
    Returns:
        int: c - новое значение счётчика, увеличенное на единицу

    Args:
        progress_bar (_type_): прогрессбар
        c (int): текущее значение счётчика
        negative (bool, optional): если надо, чтобы счётчик уменьшался, параметр должен быть равен True. По умолчанию равен False.

    Returns:
        tuple[Any, int]: progress_bar - изменённый прогрессбар, переданный функции; c - изменённый счётчик
    """
    c = c - 1 if negative else c + 1
    progress_bar.UpdateBar(c)
    return progress_bar, c

def xlsx_file_beautifulication(file_name:str|BytesIO, df_name) -> (BytesIO | None):
    """Функция делает экселевский файл "красивым" - пытается подогнать ширину столбцов под текст и делает везде все нужные границы.

    Args:
        file_name (str|BytesIO): имя изменяемого файла или файло-объект
        df_name (_type_): передаваемый датафрейм, который надо внести

    Returns:
        BytesIO | None: файло-объект, если он был передан - в ином случае, ничего не возвращаем
    """
    import xlsxwriter
    no_permission_error = False  # И мы тут снова задействуем ту переменную и ту логику, что будем пытаться получить доступ к файлу на случай, если я его не закрыл
    while not no_permission_error:
        try:
            writer = pd.ExcelWriter(file_name, engine='xlsxwriter')  # Дальше идут довольно абстрактные строчки, потому что я их взял из интернета. Вот вы знали, что в пандасе есть ЭксельРайтер? теперь знаете. И вот я не знал
        except PermissionError:  # Вообще сомневаюсь, что кто-то будет держать сводную таблицу открытой, но во избежание это надо сделать, потому что обычно сводную таблицу забываю закрыть я
            # sg.popup('Закройте файл!') 
            pass # Просто всплывающее окошко - назойливое, но настойчивое
        else:
            no_permission_error = True
    no_permission_error = False  # Ещё раз перезадаю переменную, а то мало ли

    df_name.to_excel(writer, sheet_name='sheetName', index=False, na_rep='NaN')  # Ну, тут я верю интернету наслово. Если оно и в самом деле сделает то, что надо, то ладно. Тут мы, типа, передаём значения в сам эксель (мне казалось, это задаётся немного иначе, ну ладно)
    worksheet = writer.sheets['sheetName']
    for column in df_name:  # А дальше логика такая - находим максимальную ширину, которая либо в текущей строке, либо в шапке. и перезадаём ширину всей таблички. Я верю интернету наслово - оно реально работает
        column_length = max(df_name[column].astype(str).map(len).max(), len(column))
        col_idx = df_name.columns.get_loc(column)
        worksheet.set_column(col_idx, col_idx, column_length)
    # workbook = writer.book
    # border_fmt = workbook.add_format({'bottom':1, 'top':1, 'left':1, 'right':1})
    # worksheet.conditional_format(xlsxwriter.utility.xl_range(0, 0, len(session_state['pivot_table']), len(session_state['pivot_table'].columns)), {'type': 'no_errors', 'format': border_fmt})
    worksheet.conditional_format(xlsxwriter.utility.xl_range(0, 0, len(df_name), len(df_name.columns)), {'type': 'no_errors', 'format': writer.book.add_format({'bottom':1, 'top':1, 'left':1, 'right':1})})
    writer.close()  # В старой версии Пандаса тут надо было писать сейв, а в новой - клоуз
    return file_name if isinstance(file_name, BytesIO) else None

def find_all_systems(main_system_name:str) -> list:
    """Определяет количество систем, записанных внутри одного обозначения. Так как обозначения могут быть почти что любыми, функция будет постоянно дорабатываться

    Args:
        main_system_name (str): Основное обозначение системы, из которого будут извлекаться все прочие

    Returns:
        list: all_system_names - все названия систем
    """

    system_names = main_system_name.split(',')  # Самая важная часть - разные системы могут быть перечислены через запятую. Но это не всегда так - чтобы не писать подряд идущие номера систем, их пишут через дефис (зачастую без пробелов). Так что, нужно дальше сплитать по дефису? Ни-хе-ра - в самом номере системы может содержаться дефис, и тогда что прикажете делать?
    PrintDebugModeOn(system_names)
    all_system_names = []  # На последний вопрос есть ответ, но пока нам нужен перечень ВСЕХ систем, так что создаём пустой список
    for system_name in system_names:  # Пробегаемся по всему, что разделено запятыми
        system_name = system_name.strip()
        is_match = re.fullmatch(r'(([A-Za-zА-Яа-яЁё])\d{1,})\s*?-\s*?(\2\d{1,})', system_name)  # А теперь кое-что сложное. Номера систем могут быть записаны как угодно, одному богу известно, чем руководствуются люди, записывая их. И я научился отслеживать все более-менее сложные случаи. Но самый банальный, когда у нас записано, например, "П1-П4", программа не отслеживала. Поэтому я допёр до использования дополнительной функции из ре-модуля, которая проверяет на соответствие описанной схеме
        if is_match:  # И если да...
            system_name_all = is_match.group(1, 3)  # Забираем первую и третью группы - ну, так работает этот модуль
        else:  # В ином случае...
            is_match = re.fullmatch(r'(.+?)-(\d{1,})-(\d{1,}).*', system_name)
            if is_match:
                PrintDebugModeOn(is_match.group(1, 2, 3))
                return [f"{is_match.group(1)}-{i + 1}" for i in range(int(is_match.group(2)), 1 + int(is_match.group(3)))]
            all_masks = (r'\d?[A-Za-zА-Яа-яЁё]\D?\D?\D?\D?\S?\S?\S?\S?\d{1,}[А-Яа-яЁё]*', r'\d{9}[а-яё]?')  # У нас есть набор масок, и нам их надо будет перебрать. Я буду честен - я сам не понимаю, как работает первое регулярное выражение, но оно просто работает и почти безошибочно находит все номера, причём умея отличать, где заканчивается номер в тех случаях, когда и в номере дефис, и между номерами дефис. Полагаю, что если что-то пойдёт не так и не найдётся, надо будет пошаманить ещё с выражением, но пока в этом нет нужды
            system_name_all = max((re.findall(the_mask, system_name, re.IGNORECASE) for the_mask in all_masks), key=lambda x:len(x))  # И нам нужно найти то, в котором будет найден максимум. Раньше, справедливости ради, тут было три варианта масок, но одна из них толком не работала, поэтому от неё отказался.
        system_name_all = system_name_all if system_name_all else [main_system_name]  # Иногда бывает по-идиотски записанная система, где нет цифр в конце, так что мы просто считаем, что в таком случае всего одна система и успокаиваемся
        PrintDebugModeOn(system_name_all)
        if len(system_name_all) > 1:  # А вот если у нас нашлось больше одной, это значит, что у нас записаные через дефис системы. которых дофига. И тут могут возникнуть ещё сложности. Главная сложность - как оказалось, через дефис может быть не только формата "П1.Х-П1.У", но и формата "ПХ.1-ПУ.1". Вообще, как я понял, сильнее всего отличаться будет некое число, которое и обозначает начало и конец последовательности. Поэтому нам надо вообще понять, что с этими числами не так по ходу дела
            all_positions_in_system_name = tuple(tuple(filter(None, re.findall(r'(\D?|\d+)', sys_name, re.IGNORECASE))) for sys_name in system_name_all)  # Итак, логика в чём? У нас есть два класса элементов - не-числа и числа. Не-числа всегда одинаковые. Числа могут и отличаться. Нам важно понять, что есть не-число, а что - число, и мы разбиваем всю строку на такие вот элементы
            PrintDebugModeOn(all_positions_in_system_name)
            system_condition = lambda x: x[0] == x[1]  # Потом нашей задачей будет сравнить каждый из элементов начала и конца. Сравнивать будем по этой мини-функции. Логика в том, что отличаться должно некое число, но оно может быть как в конце, так и в середине. И поэтому нам понадобится три разных переменных, из которых средняя будет разбита на две переменные для удобства. Итак
            before_changing_part = ''.join(el[0] for el in it.takewhile(system_condition, zip(all_positions_in_system_name[0], all_positions_in_system_name[1])))  # Тейквайл берёт последовательность до момента, когда нарушается условие. Условие - одинаковость элементов начального и конечного обозначения, которые мы зазипали. И чтобы два раза не вставать ещё и объединим в строчку
            after_changing_part = it.dropwhile(system_condition, zip(all_positions_in_system_name[0], all_positions_in_system_name[1]))  # Дропвайл работает диаметрально противоположно тейквайлу - забирает все значения С МОМЕНТА, когда нарушается условие. В том числе и то место, где условие нарушается. Так что я решил воспользоваться этим. Мы создаём генератор, в котором нам нужен будет первый элемент, а остальные надо будет загнать в строку. Так что...
            changing_part_start, changing_part_thend = next(after_changing_part)  # Мы для обозначения начала и конца сдвигаем генератор, распаковывая кортеж (а первый элемент в генераторе у нас кортеж и есть). Раньше для задания этой функции использовался Фильтерфолс
            after_changing_part = ''.join(el[0] for el in after_changing_part)  # А оставшееся соединяем. Генератор исчерпан, радуемся жизни и восхищаемся великолепием данной оптимизации
            PrintDebugModeOn(before_changing_part, changing_part_start, changing_part_thend, after_changing_part)
            PrintDebugModeOn(changing_part_start.isdigit(), changing_part_thend.isdigit())
            if changing_part_start.isdigit() and changing_part_thend.isdigit():  # Внезапно всплыла проблема, что порой название системы может быть написано через пробел. Не беда - собственно говоря, вся эта херобористика заточена на то, что у нас в изменяемой части будет два числа. Так что логично, что мой великолепный алгоритм будет работать тогда и только тогда, когда оба этих параметра - числа
                all_system_names += [f"{before_changing_part}{'0' * (min(len(changing_part_start), len(changing_part_thend)) - len(str(i)))}{i}{after_changing_part}" for i in range(int(changing_part_start), int(changing_part_thend) + 1)]  # А потом мы создаём все промежуточные обозначения систем. Формула выглядит очень мудрёной, но, думаю, в ней можно разобраться: соединяем начало, сколько-то нулей (пока нулей, но допускаю, что могут быть и иные случаи), число и конец.
            else:  # А если это у нас не числа, то уходим отсюда, нам тут делать нечего
                all_system_names = system_names
                break
            PrintDebugModeOn(all_system_names)
        else:
            all_system_names += system_name_all  # Ну а если у нас одна система, то её и добавляем
    return all_system_names

class Blank:
    """Общий класс для данных, которые извлекаются из бланка-заказа. В настоящий момент поддерживает бланки на ВЕРОСА, канальное оборудование, осевые вентиляторы, индустриальные вентиляторы. Чем новее бланк, тем лучше. Бланк ОБЯЗАТЕЛЬНО должен быть в форматe docx, doc, rtf, pdf!!! При попытке загнать файл другого формата будет ошибка, так что будьте внимательны. Если же файл формата doc, rtf, pdf, то могут возникнуть непредвиденные ошибки, будьте осторожны!
    В качестве аргумента принимается имя файла с бланком.
    """
    ready_blank_types = ('ВЕРОСА', 'Общепромышленные', 'Канальное оборудование', 'Индустриальный вентилятор', 'Индивидуальный тепловой пункт', 'Другое')
    all_columns = (
            'Бланк-заказ', 'Дата бланк-заказа', 'Входящий номер', 'Дата входящего номера', 'Объект', 'Номер объекта', 'Дата', 'Организация', 'Менеджер', 'Выполнил', 'Поток', 'Название', 'Типоразмер',  # Колонки основной информации
            'Назначение', 'Название блока', 'Тип блока', 'Информация о блоке'  # Колонки информации с содержимым бланка
        )

    def __init__(self, inputed:str|docx.document.Document) -> None:
        """Инициализация бланка

        Args:
            inputed (str | docx.document.Document): переданный параметр - либо адрес файла, либо документ-объект

        Raises:
            TypeError: возникает, если передано что-то не строковое и не объектно-документное, либо если файл не того формата. Возникновени других ошибок означает непредусмотренные ошибки
        """

        if isinstance(inputed, str):
            self.filename = inputed
            self.extention = os.path.splitext(inputed)[1]
            if self.extention in SUPPORTED_EXCTENTIONS_FOR_BLANK:
                if self.extention == '.docx':
                    self.Document_object, self.ordered_content, self.ordered_images = ordered_content_from_docx(self.filename)
                    self.temporary_filename = inputed
                else:
                    self.temporary_filename = temporary_filename(self.filename, new_extension='.docx')
                    convert_using_win32(self.filename, self.temporary_filename, 'docx', False)
                    self.Document_object, self.ordered_content, self.ordered_images = ordered_content_from_docx(self.temporary_filename)
            else:
                raise TypeError("Введённый файл не является файлом поддерживаемых расширений (" + ','.join(SUPPORTED_EXCTENTIONS_FOR_BLANK) + ")!")
        elif isinstance(inputed, docx.document.Document):
            self.filename, self.temporary_filename = None, None
            self.Document_object, self.ordered_content, self.ordered_images = ordered_content_from_docx(inputed)
        else:
            raise TypeError("Переданный аргумент некорректного типа!")

        bio = BytesIO()
        self.Document_object.save(bio)
        self.docx2text_objext = process(bio)

        self.blank_type = {key: False for key in self.ready_blank_types}
        self.main_information, self.all_avaiable_information = self.blank_processing(self)
        self.ALL_MAIN_INFO = []
        for zip_data in zip((data['Название блока'] for data in self.all_avaiable_information), (data['Информация о блоке'] for data in self.all_avaiable_information)):
            if zip_data not in self.ALL_MAIN_INFO:
                self.ALL_MAIN_INFO.append(zip_data)
        self.info_text = tabulate(self.main_information.items(), ('Заголовок', 'Значение'))
        self.all_main_info_text = tabulate(self.ALL_MAIN_INFO, ('Заголовок', 'Значение'))

    def __str__(self) -> str:
        return next(key for key, value in self.blank_type.items() if value) + '\n\n' + self.info_text

    # def __repr__(self) -> str:
    #     return f""

    @staticmethod
    def blank_processing(self):
        """Основная функция обработки бланков. Склеена из кусков разной степени древности, где-то оно модернизировано, где-то - нет. Но оно работает практически стабильно. Разбирает бланк почти по винтику и вытаскивает всю основную информацию (название бланк-заказа, дата, входящий номер, дата входящего номера, объект, номер объекта, дата, организация, менеджер, испиолнитель, название, типоразмер) и все параметры установки/оборудования/вентилятора.

        Returns:
            dict[str, str], list: _description_
        """

        def all_information_for_venti_or_kanal(self, start_criterion):
            """Вспомогательная функция, необходимая в тех случаях, если бланк относится к вентилятору или канальному оборудованию

            Args:
                doc (Document): объект, являющий собой выгрузку файла в программу
                start_criterion (_type_): критерий старта начала параметров (к счастью, шаблон бланков там прост, и это выявить в разы проще)

            Returns:
                list: список всяческой разной ерунды, что обычно и содержит описание
            """

            start_of_parametrs = False
            all_inform = []
            for table in self.Document_object.tables:
                for row in table.rows:
                    try:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if start_criterion in cell_text and cell_text not in all_inform and not start_of_parametrs:
                                start_of_parametrs = True
                                all_inform.append(cell_text)
                                pass

                            if start_of_parametrs and (cell_text not in all_inform) and (cell_text != ''):
                                pass
                                all_inform.append(cell_text)
                    except IndexError:
                        print('Я не знаю, с чем это связано, право слово')
                        pass
            return all_inform

        count_kanal = 0
        for table in self.Document_object.tables:
            if any(self.blank_type.values()):
                break
            for row in table.rows:
                if any(self.blank_type.values()):
                    break
                try:
                    for cell in row.cells:
                        if any(self.blank_type.values()):
                            break
                        cell_text = cell.text.strip()
                        PrintDebugModeOn(cell_text)
                        pass
                        if 'ВЕРОСА' in cell_text or 'Airmate' in cell_text:
                            self.blank_type[self.ready_blank_types[0]] = True

                        elif 'ОБЩЕПРОМЫШЛЕННЫЕ И СПЕЦИАЛЬНЫЕ ВЕНТИЛЯТОРЫ ВЕЗА' in cell_text:
                            self.blank_type[self.ready_blank_types[1]] = True

                        elif cell_text in ['ООО «ВЕЗА»', '111397, Москва, Зеленый пр-т, д20, 6 этаж', 'Тел: +7(495)989-47-20; Факс: +7(495)626-99-02', 'veza@veza.ru'] or cell_text in ['ООО "Веза"', 'Москва, Зеленый проспект д.20', 'Тел: +7 (495) 989-47-20; Факс: +7 (495) 989-47-20', 'msk1@veza.ru']:
                            count_kanal += 1
                        if count_kanal >= 4:
                            self.blank_type[self.ready_blank_types[2]] = True

                        if cell_text == 'Технические характеристики на стандартный индустриальный вентилятор':
                            self.blank_type[self.ready_blank_types[3]] = True

                        if 'Пункт тепловой индивидуальный' in cell_text:
                            self.blank_type[self.ready_blank_types[4]] = True
                except Exception as err:
                    print(err)
                    break
        for paragraph in self.Document_object.paragraphs:
            if any(self.blank_type.values()):
                    break
            if paragraph.text == 'Технические характеристики на стандартный индустриальный вентилятор':
                self.blank_type[self.ready_blank_types[3]] = True
            if 'Пункт тепловой индивидуальный' in paragraph.text:
                self.blank_type[self.ready_blank_types[4]] = True

        self.blank_type[self.ready_blank_types[5]] = not any(self.blank_type.values())
        PrintDebugModeOn(self.blank_type)
        pass

        result = {key: '-' for key in self.all_columns[0:-4]}
        ALL_RESULTS = []
        clean_information = {key: '-' for key in self.all_columns[-4:]}

        if not self.blank_type[self.ready_blank_types[5]]:
            if self.blank_type[self.ready_blank_types[0]] or self.blank_type[self.ready_blank_types[1]]:
                for paragraph in self.Document_object.paragraphs:
                    if all(value != '-' for value in result.values()):
                        break
                    paragraph_text = paragraph.text
                    if 'БЛАНК' in paragraph_text.upper():
                        result['Бланк-заказ'], result['Дата бланк-заказа'] = re.findall(r'БЛАНК[\-\s]ЗАКАЗ\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', paragraph_text, re.IGNORECASE)[0]
                    if 'входящий: ' in paragraph_text.lower():
                        result['Входящий номер'], result['Дата входящего номера'] = re.findall(r'входящий:\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', paragraph_text, re.IGNORECASE)[0]
                for table in self.Document_object.tables:
                    if all(value != '-' for value in result.values()):
                        break
                    for row in table.rows:
                        if all(value != '-' for value in result.values()):
                            break
                        try:                        
                            for cell in row.cells:
                                cell_text = cell.text
                                PrintDebugModeOn(cell_text)
                                pass    
                                if all(value != '-' for value in result.values()):
                                    break
                                PrintDebugModeOn(cell_text)

                                if 'БЛАНК' in cell_text.upper():
                                    result['Бланк-заказ'], result['Дата бланк-заказа'] = re.findall(r'БЛАНК[\-\s]ЗАКАЗ\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', cell_text, re.IGNORECASE)[0]

                                if 'входящий:' in cell_text.lower():
                                    result['Входящий номер'], result['Дата входящего номера'] = re.findall(r'входящий:\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', cell_text, re.IGNORECASE)[0]

                                if 'ОБЪЕКТ:' in cell_text.upper():
                                    result['Объект'] = re.findall(r'объект:\s?(.+)', cell_text, re.IGNORECASE)[0]
                                    if '(' in cell_text:
                                        result['Номер объекта'] = re.findall(r'\((.+?)\)', cell_text, re.IGNORECASE)[-1]

                                if 'код:' in cell_text.lower():
                                    result['Типоразмер'] = re.findall(r'код:\s?(.+)', cell_text, re.IGNORECASE)[0]

                                for column_key in self.all_columns[6:-3]:
                                    if column_key.lower() + ':' in cell_text.lower():
                                        PrintDebugModeOn(self.all_columns, self.all_columns[6:-3], cell_text)
                                        result[column_key] = re.findall(fr"{column_key}:\s?(.+)", cell_text, re.IGNORECASE)[0]
                                pass

                        except Exception as err:
                            print(err)
                            break
                for result_key in result:
                    # PrintDebugModeOn(result_key, result_file[result_key])
                    abzatc = result[result_key].split('\n')
                    PrintDebugModeOn(abzatc)
                    if len(abzatc) > 1:
                        if result_key.upper() in result[result_key].upper():
                            for strochka in abzatc:
                                PrintDebugModeOn(strochka)
                                if result_key.upper() in strochka.upper():
                                    result[result_key] = strochka.split(':')[1].lstrip()
                        else:
                            result[result_key] = abzatc[0].lstrip()

                if self.blank_type[self.ready_blank_types[0]]:
                    PrintDebugModeOn(*self.ordered_content, sep='\n')
                    table_count = sum(1 for item in self.ordered_content if item[0] != '' and isinstance(item[1], int))
                    other_count = sum(1 for item in self.ordered_content if item[0] != '' and not isinstance(item[1], int))
                    PrintDebugModeOn(table_count, other_count)
                    mostly_tables = table_count > other_count
                    PrintDebugModeOn(mostly_tables)

                    for i in range(len(self.ordered_content)):
                        if isinstance(self.ordered_content[i][1], int):
                            table_content = []
                            tek_table = self.ordered_content[i][0]
                            if mostly_tables:
                                for row in tek_table.rows:
                                    for cell in row.cells:
                                        table_content.append(cell.text)
                            else:
                                for column in tek_table.columns:
                                    for cell in column.cells:
                                        table_content.append(cell.text)
                            PrintDebugModeOn(table_content)
                            self.ordered_content[i][0] = table_content
                    PrintDebugModeOn(*self.ordered_content, sep='\n')
                    pass

                    only_content = [item[0] for item in self.ordered_content if item[0] != '' or ('стр ' not in item[0] and ' / ' not in item[0]) or ('kck' not in item[0].lower()) or ('бланк заказ' not in cell.lower())]
                    # type_content = [type(item[0]) for item in self.ordered_content if item[0] != '']
                    PrintDebugModeOn(*only_content, sep='\n')
                    pass

                    monoblocks, blocks = 0, 0
                    # mono_flag, bloc_flag = False, False
                    slash_checked = False
                    for table in self.Document_object.tables:
                        if monoblocks and blocks:
                            break
                        for row in table.rows:
                            if monoblocks and blocks:
                                break
                            try:
                                for cell in row.cells:
                                    cell_text = cell.text
                                    if 'моноблоков' in cell_text.lower() and not monoblocks:
                                        PrintDebugModeOn(cell_text)
                                        if '/' in cell_text and not slash_checked:
                                            slash_checked = True
                                            try:
                                                chisla = cell_text.split('=')[1][:-2]
                                            except:
                                                chisla = cell_text.split('\n')
                                                PrintDebugModeOn(chisla)
                                                for chi in chisla:
                                                    PrintDebugModeOn(chi)
                                                    if 'моноблоков' in chi.lower():
                                                        chisla = chi.split()[1]
                                                        PrintDebugModeOn(chisla)
                                                        break
                                            finally:
                                                PrintDebugModeOn(chisla)
                                                chisla = chisla.split('/')
                                                PrintDebugModeOn(chisla)
                                                monoblocks, blocks = chisla[1], chisla[0]
                                                # mono_flag, bloc_flag = True, True
                                                
                                        else:
                                            if not monoblocks:
                                                monoblocks = cell_text.split('=')[1][:-2]
                                                # mono_flag = True
                                                


                                    if 'блоков' in cell_text.lower() and 'моноблоков' not in cell_text.lower() and not blocks:
                                        PrintDebugModeOn(cell_text)
                                        if '/' in cell_text and not slash_checked:
                                            slash_checked = True
                                            try:
                                                chisla = cell_text.split('=')[1][:-2]
                                            except:
                                                chisla = cell_text.split('\n')
                                                PrintDebugModeOn(chisla)
                                                for chi in chisla:
                                                    PrintDebugModeOn(chi)
                                                    if 'моноблоков' in chi.lower():
                                                        chisla = chi.split()[1]
                                                        PrintDebugModeOn(chisla)
                                                        break
                                            finally:
                                                PrintDebugModeOn(chisla)
                                                chisla = chisla.split('/')
                                                PrintDebugModeOn(chisla)
                                                monoblocks, blocks = chisla[1], chisla[0]
                                                # mono_flag, bloc_flag = True, True
                                                
                                        else:
                                            if not blocks:
                                                blocks = cell_text.split('=')[1][:-2]
                                                # bloc_flag = True
                                                
                            except Exception as err:
                                print(err)
                                break
                    monoblocks, blocks = int(monoblocks), int(blocks)

                    all_needed_information = []
                    for i in range(1, monoblocks + 1):
                        dob_mono = str(i) + '.'
                        needed_information = {
                            'value' : dob_mono,
                            'type' : 'моноблок',
                            'title' : '',
                            'found_value' : False,
                            'information' : [],
                            'found_info' : False}
                        all_needed_information.append(needed_information)
                        for j in range(1, blocks + 1 + 1):
                            dob_bloc = dob_mono + str(j) + '.'
                            needed_information = {
                                'value' : dob_bloc,
                                'type' : 'блок',
                                'title' : '',
                                'found_value' : False,
                                'information' : [],
                                'found_info' : False}
                            all_needed_information.append(needed_information)
                    # checking_arguments(all_needed_information)
                    all_headers = [need_info['value'] for need_info in all_needed_information]
                    PrintDebugModeOn(all_headers)
                    pass

                    for i in range(len(all_needed_information) - 1):
                        prev_abz = ''
                        PrintDebugModeOn(all_needed_information[i]['value'], all_needed_information[i]['found_value'])
                        pass
                        if not mostly_tables:
                            for content in only_content:
                                if all_needed_information[i]['found_value'] and (prev_abz != '') and (prev_abz != content) and (content != '')  and not all_needed_information[i]['found_info']:
                                    # checking_arguments(content)
                                    first_content = content.split()[0] if isinstance(content, str) else content[0]
                                    if any(left_header in first_content for left_header in all_headers[i + 1:]):
                                        PrintDebugModeOn('Что, выходим?', content)
                                        PrintDebugModeOn(all_needed_information[i]['information'])
                                        all_needed_information[i]['found_info'] = True
                                        all_needed_information[i]['information'] = '\n'.join(all_needed_information[i]['information'])
                                        pass
                                        break
                                    else:
                                        PrintDebugModeOn('Нужный ли это абзац:', content)
                                        if isinstance(content, str):
                                            all_needed_information[i]['information'].append(content)
                                        else:
                                            all_needed_information[i]['information'].append('; '.join(content))
                                        pass
                                        if 'Спектральные и суммарные уровни звуковой мощности' in content or 'Автоматика' in content:
                                            if all_needed_information[i]['information']:
                                                all_needed_information[i]['information'] = '\n'.join(all_needed_information[i]['information'][:-1])
                                            all_needed_information[i]['found_info'] = True
                                            break

                                if all_needed_information[i]['value'] in content and (all_needed_information[i]['value'] == content[:len(all_needed_information[i]['value'])]) and not all_needed_information[i]['found_value'] and (content != prev_abz):
                                    PrintDebugModeOn('Абзац:', content)
                                    all_needed_information[i]['title'] = content
                                    all_needed_information[i]['found_value'] = True
                                    prev_abz = content

                                if ('Должность,ФИО,подпись' in content):
                                    PrintDebugModeOn('Что, выходим?', content)
                                    PrintDebugModeOn(all_needed_information[i]['information'])
                                    if all_needed_information[i]['information']:
                                        all_needed_information[i]['information'] = '\n'.join(all_needed_information[i]['information'][:-2])
                                        PrintDebugModeOn(all_needed_information[i]['information'])
                                        pass
                                    all_needed_information[i]['found_info'] = True
                                    break

                        else:
                            for content in only_content:
                                if isinstance(content, list):
                                    for cell in content:
                                        PrintDebugModeOn(cell)
                                        # pass
                                        if all_needed_information[i]['found_value'] and (prev_abz != '') and (prev_abz != cell) and (cell != '')  and not all_needed_information[i]['found_info'] and ('стр ' not in cell and ' / ' not in cell) and ('kck' not in cell.lower()) and ('бланк заказ' not in cell.lower()):
                                            PrintDebugModeOn('Нужный ли это абзац:', cell)
                                            # pass
                                            all_needed_information[i]['information'] = cell
                                            all_needed_information[i]['found_info'] = True
                                            # pass
                                            break

                                        if all_needed_information[i]['value'] in cell and (all_needed_information[i]['value'] == cell[:len(all_needed_information[i]['value'])]) and not all_needed_information[i]['found_value'] and (cell != prev_abz) and ('стр ' not in cell and ' / ' not in cell) and ('kck' not in cell.lower()) and ('бланк заказ' not in cell.lower()):
                                            PrintDebugModeOn('Абзац:', cell)
                                            # pass
                                            all_needed_information[i]['title'] = cell
                                            all_needed_information[i]['found_value'] = True
                                            prev_abz = cell
                    PrintDebugModeOn(*all_needed_information, sep='\n')
                    # pass

                    for additional_information in all_needed_information:
                        if (additional_information['found_value'] and additional_information['found_info']) or (additional_information['title'] and additional_information['information']):
                            clean_information = {
                                'Название блока' : additional_information['title'],
                                'Тип блока' : additional_information['type'],
                                'Информация о блоке' : additional_information['information']}
                            PrintDebugModeOn(clean_information)
                            itog_information = result | clean_information
                            ALL_RESULTS.append(itog_information)
                else:
                    thend_criterion = 'Спектральные уровни звуковой мощности'
                    all_information = all_information_for_venti_or_kanal(self, result['Типоразмер'])        
                    clean_information = {
                        'Название блока' : all_information[0],
                        'Тип блока' : 'моноблок',
                        'Информация о блоке' : '; '.join(el.strip() for el in (all_information[1:all_information.index(thend_criterion)] + [item for item in all_information if 'лнительн' in item])).replace('\n', '; ')}
                    itog_information = result | clean_information
                    ALL_RESULTS.append(itog_information)
                    pass
                pass

            elif self.blank_type[self.ready_blank_types[2]]:
                indexes = []
                for table in self.Document_object.tables:
                    if all(value != '-' for value in result.values()):
                        break
                    for row in table.rows:
                        flag_next_cell = False
                        text_prev_cell = ''
                        PrintDebugModeOn(flag_next_cell, text_prev_cell)
                        if all(value != '-' for value in result.values()):
                            break
                        try:
                            for cell in row.cells:
                                if all(value != '-' for value in result.values()):
                                    break
                                cell_text = cell.text
                                PrintDebugModeOn(cell_text)
                                prov = cell_text.upper()

                                if 'ПРОЕКТ' in prov:
                                    result['Входящий номер'] = cell_text.split()[1]

                                if flag_next_cell and text_prev_cell != cell_text:
                                    PrintDebugModeOn(cell_text)
                                    match text_prev_cell:
                                        case 'Объект:':
                                            result['Объект'] = cell_text
                                        case 'Заказчик:':
                                            result['Организация'] = cell_text
                                        case 'Исполнитель:':
                                            result['Выполнил'] = cell_text
                                        case 'Название:':
                                            result['Бланк-заказ'] = cell_text
                                            result['Название'] = cell_text
                                    
                                    flag_next_cell = False

                                if cell_text in ['Объект:', 'Заказчик:', 'Исполнитель:', 'Название:']:
                                    flag_next_cell = True
                                    text_prev_cell = cell_text

                                if 'Индекс:' in cell_text:
                                    PrintDebugModeOn(cell_text)
                                    if cell_text.split()[1] not in indexes:
                                        indexes.append(cell_text.split()[1])

                        except Exception as err:
                            print(err)
                            break
                PrintDebugModeOn(indexes)
                result['Типоразмер'] = '; '.join(indexes)

                thend_criterion = 'Спектральные (дБ) и суммарные (дБА) уровни звуковой мощности'
                all_information = all_information_for_venti_or_kanal(self, 'Характеристики входящего оборудования')
                PrintDebugModeOn(all_information)
                # all_information = all_information[1:all_information.index(thend_criterion)] + all_information[all_information.index('Дополнительное оборудование:'):all_information.index('Габаритная схема')]
                all_needed_information = []
                all_main_blocks = [piece_of_info for piece_of_info in all_information if piece_of_info.split()[0][-1] == '.']
                main_block_name = [' '.join(_.split()[1:]) for _ in all_main_blocks]
                main_block_info = []
                PrintDebugModeOn(all_main_blocks, main_block_name)
                pass
                for i in range(len(all_main_blocks)):
                    clean_information['Название блока'] = all_main_blocks[i]

                    tek_main_block_name = ' '.join(all_main_blocks[i].split()[1:])
                    if all_main_blocks[i] != all_main_blocks[-1]:
                        finis_criterion = all_main_blocks[i + 1]
                    else:
                        finis_criterion = thend_criterion

                    if tek_main_block_name not in main_block_name[:i]:
                        try:
                            clean_information['Информация о блоке'] = all_information[all_information.index(all_main_blocks[i]) + 1:all_information.index(finis_criterion)]
                        except ValueError:
                            finis_criterion = 'Корректированный уровень звукового давления LpA, дБ(А)'
                            clean_information['Информация о блоке'] = all_information[all_information.index(all_main_blocks[i]) + 1:all_information.index(finis_criterion)]
                        main_block_info.append(all_information[all_information.index(all_main_blocks[i]) + 1:all_information.index(finis_criterion)])
                    else:
                        clean_information['Информация о блоке'] = main_block_info[main_block_name.index(tek_main_block_name)]
                        main_block_info.append(main_block_info[main_block_name.index(tek_main_block_name)])
                    
                    # checking_arguments(clean_information['Информация о блоке'])
                    clean_information['Информация о блоке'] = '; '.join(clean_information['Информация о блоке'])
                    all_needed_information.append(clean_information)
                    
                    itog_information = result | clean_information
                    # checking_arguments(itog_information)
                    ALL_RESULTS.append(itog_information)
                if 'Дополнительное оборудование:' in all_information and 'Габаритная схема' in all_information:
                    clean_information['Название блока'] = 'Дополнительное оборудование'
                    clean_information['Информация о блоке'] = '; '.join(all_information[all_information.index('Дополнительное оборудование:') + 1:all_information.index('Габаритная схема')])
                    itog_information = result | clean_information
                    ALL_RESULTS.append(itog_information)
                PrintDebugModeOn(ALL_RESULTS)
                pass

            elif self.blank_type[self.ready_blank_types[3]]:
                temp_keys = ('Проект', 'Дата бланк-заказа', 'Входящий номер', 'Дата входящего номера', 'Объект', 'Сисиема ', 'Дата', 'Заказчик', 'Менеджер', 'Исполнитель', 'Система', 'Название', 'Модель:', 'Вентилятор:')
                temp_stroka = {key: '-' for key in temp_keys}

                activated = ''  # Эта переменная означает, какой именно ключ должен быть активирован на следующем проходе
                next_cell = False  # Эта переменная означает, что в следующую ячейку надо внести значения
                all_points = {key : False for key in temp_stroka.keys()}  # 
                
                # for table in doc.tables[1:5]:
                for table in self.Document_object.tables[0:1]:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text  # Просто чтобы не писать это раз за разом
                            if next_cell and cell_text != activated:
                                PrintDebugModeOn(cell_text)
                                result[list(result.keys())[temp_keys.index(activated)]] = cell_text
                                pass
                                next_cell = False

                            if cell_text in all_points and not all_points[cell_text]:  # Итак, если мы нашли один из ключей, и при этом его ещё не активировали, то...
                                PrintDebugModeOn(cell_text)
                                activated = cell_text  # ...сохраняем активированный ключ, чтобы потом (выше по тексту) по этому ключу внести значения
                                all_points[cell_text] = True  # Даём понять, какой именно ключ активирован
                                next_cell = True  # Даём понять, что следующая ячейка - нужная
                
                all_information = []
                # for row in self.Document_object.tables[1].rows[3:-1]:
                for row in self.Document_object.tables[0].rows[3:6]:
                    # for cell in row.cells[2:]:
                    for cell in row.cells[5:]:
                        PrintDebugModeOn(cell.text)
                        if cell.text not in all_information:
                            all_information.append(cell.text)
                PrintDebugModeOn(all_information)
                # for row in self.Document_object.tables[2].rows[2:]:
                for row in self.Document_object.tables[0].rows[7:]:
                    for cell in row.cells:
                        if cell.text not in all_information and cell.text:
                            all_information.append(cell.text)
                PrintDebugModeOn(all_information)
                clean_information = {
                    'Название блока' : result['Типоразмер'],
                    'Тип блока' : 'моноблок',
                    'Информация о блоке' : '; '.join(all_information)
                }
                itog_information = result | clean_information
                ALL_RESULTS.append(itog_information)
                add_information = []
                # for row in self.Document_object.tables[4].rows[1:]:
                for row in self.Document_object.tables[1].rows[5:11]:
                    for cell in row.cells:
                        if cell.text not in add_information and cell.text:
                            add_information.append(cell.text)
                PrintDebugModeOn(add_information)
                clean_information = {
                    'Название блока' : 'Примечания: Комплектация и технические особенности(отдельно по счёту)', #self.Document_object.tables[1].rows[4].cells[0].text, #self.Document_object.tables[4].rows[0].cells[0].text,
                    'Тип блока' : 'моноблок',
                    'Информация о блоке' : '; '.join(add_information)
                }
                itog_information = result | clean_information
                ALL_RESULTS.append(itog_information)
                PrintDebugModeOn(ALL_RESULTS)
                pass

            elif self.blank_type[self.ready_blank_types[4]]:
                print(self.docx2text_objext)
                print(*(content for content in self.ordered_content if content[0]), sep='\n')
                pass
                print([[cell.text for cell in row.cells][-2:] for i in range(3, 5) for row in self.Document_object.tables[i].rows])
                pass

                for paragraph in self.Document_object.paragraphs:
                    print(paragraph.text)
                    pass

                for table in self.Document_object.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            print(cell.text)
                            if 'Блок клапанный Физтех БКН2-47' in cell.text:
                                pass
                            pass
                        pass
                    pass
                pass

                result['Типоразмер'], result['Бланк-заказ'], result['Дата бланк-заказа'] = re.findall(r'(.+_)\s?(.+) от (.+)г', cell_text.split('\n')[1], re.IGNORECASE)[0]
                result['Организация'], result['Выполнил'], result['Объект'], result['Менеджер'] = (re.findall(fr'{key}:\s?(.+)', self.docx2text_objext, re.IGNORECASE)[0] for key in ('Заказчик', 'Выполнил', 'Объект', 'Менеджер'))
                PrintDebugModeOn(result)
                pass

                clean_information['Название блока'] = 'Вся информация о бланке'
                clean_information['Информация о блоке'] = [[cell.text for cell in row.cells][-2:] for i in range(1, 4) for row in self.Document_object.tables[i].rows]
                ALL_RESULTS.append(result | clean_information)

                # clean_information['Название блока'] = '2. Исходные данные и режим работы'
                # clean_information['Информация о блоке'] = [[cell.text for cell in row.cells] for row in self.Document_object.tables[2].rows]
                # ALL_RESULTS.append(result | clean_information)

                # clean_information['Название блока'] = '3. Основное оборудование'
                # clean_information['Информация о блоке'] = [[cell.text for cell in row.cells] for row in self.Document_object.tables[3].rows]
                # ALL_RESULTS.append(result | clean_information)

                PrintDebugModeOn(ALL_RESULTS)
                pass

        else:
            itog_information = result | clean_information
            ALL_RESULTS.append(itog_information)

        if self.filename != self.temporary_filename:
            os.remove(self.temporary_filename)
        return result, ALL_RESULTS

class DocxExpand:
    """Дополнительные методы или функции, которые должны расширить функционал модуля docx, в основном тут заимствованный у чужих код
    """

    @staticmethod
    def delete_paragraph(paragraph):
        """Удаляет выбранный параграф

        Args:
            paragraph (_type_): объект-параграф
        """
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    @staticmethod
    def add_paragraph_after_paragraph(previous_paragraph, new_paragraph_text:str, document):
        """_summary_

        Args:
            previous_paragraph (_type_): объект-параграф
            new_paragraph_text (str): текст нового параграфа
            document (_type_): объект-документ

        Returns:
            _type_: new_paragraph, нужно, чтобы иметь возможность добавить что-то после него.
        """
        new_paragraph = document.add_paragraph(new_paragraph_text)
        previous_paragraph._p.addnext(new_paragraph._p)
        return new_paragraph

    @staticmethod
    def add_paragraph_before_table(paragraph_text:str, table, document):
        """Добавляет параграф перед таблицей

        Args:
            paragraph_text (str): текст нового параграфа
            table (_type_): объект-таблица
            document (_type_): объект-документ
        """
        table._element.addprevious(document.add_paragraph(paragraph_text)._p)

    @staticmethod
    def add_paragraph_after_table(table):
        """Добавляет параграф после таблицы

        Args:
            table (_type_): объект-таблица

        Returns:
            Paragraph: объект-параграф (так надо)
        """
        return Paragraph(table._tbl.getnext(), table._parent).insert_paragraph_before()
    
    @staticmethod
    def move_table_after_paragraph(table, paragraph):
        """Передвигает таблицу после параграфа

        Args:
            table (_type_): объект-таблица
            paragraph (_type_): объект-параграф
        """
        paragraph._p.addnext(table._tbl)
    
    @staticmethod
    def remove_row(table, row):
        """Удаляет выбранную строку данной таблицы

        Args:
            table (_type_): объект-таблица
            row (_type_): объект-строка данной таблицы
        """
        table._tbl.remove(row._tr)
    
    @staticmethod
    def change_cell_applying_its_style(cell, text:str):
        """Недавно обнаруженный костыль выявил занятный факт - внутри ячеек таблицы находятся параграфы, и их можно извлечь как список параграфов, и если не передавать тексту параграфа новое значение, а просто добавить новое, то тогда сохранится исходный стиль, который был у текста в ячейке. Важно - сохраняется стиль, а если к самому стилю были применены изменения внутри шаблона, то эти изменения не сохранятся. То есть, если внутри ячейки текст стиля "Нормальный", стиль "Нормальный" - 12 кегль, во всю ширину, но в шаблоне в ячейке он сделан жирным, жирнота спадёт. Отсюда совет - можно создавать бесконечно много стилей внутри документа под каждый случай

        Args:
            cell (_type_): ячейка-объект
            text (str): новый текст
        """
        for paragraph in cell.paragraphs:
            paragraph.text += text
    pass

# ===============================================================================================================================================================================================

def from_base10_to_baseXX(number_10:int|float, base:int) -> str:
    from decimal import Decimal
    import math

    """Перевод числа из десятеричной системы счисления в заданную. Внимание - в дробных числах не гарантируется абсолютная точность!

    Args:
        number_10 (int|float): число в десятеричной системе счисления
        base (int): основание новой системы счисления
    
    Returns:
        str: число в новой системе счисления в формате str. Для дальнейшей работы с этим результатом желательно использовать в связке с классом BaseXX
    """

    def inter_and_fract_parts(number:Decimal):
        """Разделение заданного числа на целую и дробную части

        Args:
            number (Decimal): число в формате Decimal

        Returns:
            tuple[int, Decimal]: целая и дробная части
        """

        in_p = math.trunc(abs(number))
        fr_p = abs(number) - in_p
        return in_p, fr_p    

    number_10 = Decimal(str(number_10))
    sign, _, e = number_10.as_tuple()
    inter_part, fract_part = inter_and_fract_parts(number_10)

    inter_part_res = []
    if inter_part == 0:
        inter_part_res = ['0']
    else:
        while inter_part > 0:
            inter_part_res.insert(0, ALL_DIGITS[inter_part % base])
            inter_part //= base

    fract_part_res = []
    if e:
        for _ in range(-e + math.ceil(2 / math.log10(base))):
            fract_part = fract_part * base
            digit, fract_part = inter_and_fract_parts(fract_part)
            fract_part_res.append(ALL_DIGITS[digit])
    result = f"{'-' if sign else ''}{''.join(inter_part_res)}{';' + ''.join(fract_part_res) if fract_part_res else ''}"
    return result

def from_baseXX_to_base10(number_XX:str, base:int) -> int|float:
    """Перевод числа из заданной системы счисления в десятеричную. Внимание - в дробных числах не гарантируется абсолютная точность!

    Args:
        number_XX (str): число в заданной системе счисления. Передаётся как строка. Предполагается, что эту функцию используют в связке с класом BaseXX
        base (int): основание исходной системы счисления

    Returns:
        int|float: число в десятеричной системе счисления
    """
    if number_XX[0] == '-':
        sign = -1
        number_XX = number_XX[1:]
    else:
        sign = 1
    if ';' not in number_XX:
        number_XX += ';'
    inter_part, fract_part = number_XX.split(';')
    inter_part_res = sum(ALL_DIGITS.index(inter_part[-1 - i]) * pow(base, i) for i in range(len(inter_part)))
    fract_part_res = sum(ALL_DIGITS.index(fract_part[i]) * pow(base, -1 - i) for i in range(len(fract_part))) if fract_part else 0
    result = (inter_part_res + fract_part_res) * sign
    return result

class BaseXX():
    """Класс, представляющий собой число в заданной системе счисления. Поддерживает системы счисления с 2 по 60. Для систем счисления с основанием 2, 8, 10, 16 всё-таки рекомендуется использовать внутренние инструменты. Может работать в двух режимах:

    1. Если в качестве числа задано число (тип int, float), то предполагается, что результат будет число, переведённое из десятеричной системы счисления в заданную
    2. Если в качестве числа задана строка, то предполагается, что это уже число в заданной системе счисления.

    Словом, BaseXX(12, 7) и BaseXX('12', 7) будут двумя разными числами!
    """    

    def __init__(self, number:int|float|str, base:int) -> None:
        if (base < 2) or (60 < base):
            raise ValueError
        self.base = base
        self.digits = ALL_DIGITS[0:self.base]
        if isinstance(number, str):
            self.numberXX = number
            self.number10 = from_baseXX_to_base10(self.numberXX, self.base)
        else:
            self.numberXX = from_base10_to_baseXX(number, self.base)
            self.number10 = number
        pass

    def __str__(self) -> str:
        return f"{self.numberXX} [base{self.base}]"
    
    def __add__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 + other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 + Base12(other).number10)
        else:
            return BaseXX(self.number10 + other.number10 if isinstance(other, type(self)) else self.number10 + BaseXX(other, self.base).number10, self.base)
    def __radd__(self, other):
        if isinstance(self, Base12):
            return self + Base12(other)
        else:
            return self + BaseXX(other, self.base)
    def __iadd__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 + other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 + Base12(other).number10)
        else:
            return BaseXX(self.number10 + other.number10 if isinstance(other, type(self)) else self.number10 + BaseXX(other, self.base).number10, self.base)
    
    def __sub__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 - other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 - BaseXX(other).number10)
        else:
            return BaseXX(self.number10 - other.number10 if isinstance(other, type(self)) else self.number10 - BaseXX(other, self.base).number10, self.base)
    def __rsub__(self, other):
        if isinstance(self, Base12):
            return self - Base12(other)
        else:
            return self - BaseXX(other, self.base)
    def __isub__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 - other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 - BaseXX(other).number10)
        else:
            return BaseXX(self.number10 - other.number10 if isinstance(other, type(self)) else self.number10 - BaseXX(other, self.base).number10, self.base)
    
    def __mul__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 * other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 * Base12(other).number10)
        else:
            return BaseXX(self.number10 * other.number10 if isinstance(other, type(self)) else self.number10 * BaseXX(other, self.base).number10, self.base)
    def __rmul__(self, other):
        if isinstance(self, Base12):
            return self * Base12(other)
        else:
            return self * BaseXX(other, self.base)
    def __imul__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 * other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 * Base12(other).number10)
        else:
            return BaseXX(self.number10 * other.number10 if isinstance(other, type(self)) else self.number10 * BaseXX(other, self.base).number10, self.base)
    
    def __truediv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 / other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 / Base12(other).number10)
        else:
            return BaseXX(self.number10 / other.number10 if isinstance(other, type(self)) else self.number10 / BaseXX(other, self.base).number10, self.base)
    def __rtruediv__(self, other):
        if isinstance(self, Base12):
            return self / Base12(other)
        else:
            return self / BaseXX(other, self.base)
    def __itruediv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 / other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 / Base12(other).number10)
        else:
            return BaseXX(self.number10 / other.number10 if isinstance(other, type(self)) else self.number10 / BaseXX(other, self.base).number10, self.base)
    
    def __floordiv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 // other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 // Base12(other).number10)
        else:
            return BaseXX(self.number10 // other.number10 if isinstance(other, type(self)) else self.number10 // BaseXX(other, self.base).number10, self.base)
    def __rfloordiv__(self, other):
        if isinstance(self, Base12):
            return self // Base12(other)
        else:
            return self // BaseXX(other, self.base)
    def __ifloordiv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 // other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 // Base12(other).number10)
        else:
            return BaseXX(self.number10 // other.number10 if isinstance(other, type(self)) else self.number10 // BaseXX(other, self.base).number10, self.base)        

    def __mod__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 % other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 % Base12(other).number10)
        else:
            return BaseXX(self.number10 % other.number10 if isinstance(other, type(self)) else self.number10 % BaseXX(other, self.base).number10, self.base)
    def __rmod__(self, other):
        if isinstance(self, Base12):
            return self % Base12(other)
        else:
            return self % BaseXX(other, self.base)
    def __imod__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 % other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 % Base12(other).number10)
        else:
            return BaseXX(self.number10 % other.number10 if isinstance(other, type(self)) else self.number10 % BaseXX(other, self.base).number10, self.base)
    
    def __lt__(self, other):
        if isinstance(self, Base12):
            return self.number10 < other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 < Base12(other).number10
        else:
            return self.number10 < other.number10 if isinstance(other, type(self)) else self.number10 < BaseXX(other, self.base).number10
    
    def __le__(self, other):
        if isinstance(self, Base12):
            return self.number10 <= other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 <= Base12(other).number10
        else:
            return self.number10 <= other.number10 if isinstance(other, type(self)) else self.number10 <= BaseXX(other, self.base).number10
    
    def __eq__(self, other):
        if isinstance(self, Base12):
            return self.number10 == other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 == Base12(other).number10
        else:
            return self.number10 == other.number10 if isinstance(other, type(self)) else self.number10 == BaseXX(other, self.base).number10
    
    def __ne__(self, other):
        if isinstance(self, Base12):
            return self.number10 != other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 != Base12(other).number10
        else:
            return self.number10 != other.number10 if isinstance(other, type(self)) else self.number10 != BaseXX(other, self.base).number10
    
    def __gt__(self, other):
        if isinstance(self, Base12):
            return self.number10 > other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 > Base12(other).number10
        else:
            return self.number10 > other.number10 if isinstance(other, type(self)) else self.number10 > BaseXX(other, self.base).number10
    
    def __ge__(self, other):
        if isinstance(self, Base12):
            return self.number10 >= other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 >= Base12(other).number10
        else:
            return self.number10 >= other.number10 if isinstance(other, type(self)) else self.number10 >= BaseXX(other, self.base).number10

class Base12(BaseXX):
    """Класс чисел в 12-ичной системе счисления. Является наследником класса BaseXX, но уже с заданной системой счисления - 12. В качестве цифр для обозначения "10" и "11" используются специальные знаки "↊" и "↋". При задании числа в 12-ичной можно использовать стандартные "A" и "B" (всё равно ни у кого на клавиатуре не будет тех знаков)
    """
    def __init__(self, number:int|float|str) -> None:
        self.base = 12
        if isinstance(number, str) and ('↊' in number or '↋' in number):
            number = number.replace('↊', 'A').replace('↋', 'B')
        super().__init__(number, self.base)
        self.digits = ALL_DIGITS[0:10] + '↊↋'
    
    def __str__(self) -> str:
        return f"{self.numberXX} [base{self.base}]".replace('A', '↊').replace('B', '↋')

def do_something_fun():
    """Код, предложенный Везантиной на запрос "Можешь написать на Питоне код, который делает что-нибудь прикольное?". По её словам, "небольшой код на Python, который создаст анимированный градиентный экран"
    """
    import turtle
    from random import randint

    # настройки экрана
    screen = turtle.Screen()
    screen.setup(700, 700)
    screen.bgcolor("black")

    # настройки черепашки
    turtle.speed(0)
    turtle.penup()

    # создаем цветовую палитру
    colors = ["red", "orange", "yellow", "green", "blue", "purple"]

    # создаем градиентный экран
    for y in range(-350, 350, 10):
        for x in range(-350, 350, 10):
            color = colors[randint(0, len(colors)-1)]
            turtle.goto(x, y)
            turtle.dot(10, color)

    # создаем анимацию
    for _ in range(300):
        turtle.clear()
        for y in range(-350, 350, 10):
            for x in range(-350, 350, 10):
                color = colors[randint(0, len(colors)-1)]
                turtle.goto(x, y)
                turtle.dot(10, color)
            
        # поворот и перемещение черепашки
        turtle.left(5)
        turtle.forward(10)

    turtle.done()

# ===============================================================================================================================================================================================

if __name__ == '__main__':
    a = BaseXX(5, 7)
    PrintDebugModeOn(a % 2 == 0)
    a = a / 2 if a % 2 == 0 else 3*a + 1
    PrintDebugModeOn(a)
    a = Base12(18*4*3)
    print(a, a.number10)

    example_dict = {
        'key1': 'value1',
        'key2': {
            'subkey1': 'subvalue1',
            'subkey2': {
                'subsubkey1': 'subsubvalue1'
            }
        },
        'key3': 'value3'
    }
    result_dict = flatten_dictionary(example_dict, True)
    print(result_dict)
    result_dict = flatten_dictionary(example_dict, False)
    print(result_dict)

    # do_something_fun()

    PrintDebugModeOn(find_all_systems('V-FC-0001, V-FC-0003, V-FC-0005, V-FC-0011- V-FC-0015'))

    r'\d?[A-Za-zА-Яа-яЁё]\D?\D?\D?\D?\S?\S?\S?\S?\d{1,}[А-Яа-яЁё]?'

    for_automata = 'C:\\Users\\ovchinnikov\\Documents\\Python\\21_AUTOMATA\\_ТЗ_для программы\\'
    all_test_files = [
        '211012073-ОПР АПБП (Складской комплекс ООО ЛА МАРЕ г. Москва, ул. Складочная д.15).docx',  # 0
        f'{for_automata}Бланк2\\201006039-ОПР Ташир ГК(Апарт-Отель г.Ярославль).docx',  # 1
        f'{for_automata}Бланк3\\1_211043062б-ОПР ИНИЦИАТИВА ИНЖПРОЕКТ (ст.Каширская Западный Вестибюль).docx',  # 2
        'C:/Users/ovchinnikov/Documents/Python/21 AUTOMATA/_ТЗ_для программы/Бланк2/КА201006039-ОПР (с КР для завода).doc',  # 3
        f'{for_automata}Бланк9-10\\221027641д-ОПР ФАРМ ДИЗАЙН ООО (Виварий) П2С осн, В7С осн.doc',  # 4
        '221057784-ОПР П5 ГАЗПРОМ ТРАНСГАЗ Центр Диагностики и Реабилитации Воскресенкс.docx',  # 5
        '211043780б-ОПР М1 ПРОЕКТ ООО (Центр ядерной медицины г. Екатеринбург) П3.doc',  # 6
        'ДУ1-ДУ6.docx',  # 7
        'В2.6.docx',  # 8
        'В2.2, В2.4, В.20, В2.21.docx',  # 9
        'В2.1, В2.5, В2.8, В2.10, В2.12, В2.15, В2.18, В2.19, В2.22.docx',  # 10
        'В1.24.docx',  # 11
        'В1.6, В1.7.docx',  # 12
        'В1.9, В1.16.docx',  # 13
        'В1.1, В1.5, В1.8, В1.11, В1.12, В1.15, В1.18, В1.19, В1.22.docx',  # 14
        '2.docx',  # 15
        '3.docx',  # 16
        '1.docx',  # 17
        '221054633-ОПР ПВ1 ГРОСС ИНЖИНИРИНГ (Гостиница Комета. г.Москва, ул. Космонавта Волкова, 14 ).doc',  # 18
        f'{for_automata}Каналка\\ПВитп.docx',  # 19
        'П1.docx',   # 20
        'П2.docx',   # 21
        'П3.docx',   # 22
        'П4.docx',   # 23
        'ПВитп.docx',   # 24
        'V-FC-0016 - Регина.docx',  # 25
        'ДОО на 220. п. Сосенское, пос. коммунарка ППТ 2-3 участок № 32.1-П4 - Лена.rtf',  # 26
        f'{for_automata}Каналка\\П7В81 - Лена.docx',  # 27
        f'{for_automata}Веросы\\Бланк1\\211027853-ОПР АЭРОПРОЕКТ (Реконстр. аэропортового комплекса Чертовицкое г.Воронеж ОАСС).docx',  # 28
        f'{for_automata}Веросы\\221049161-ОПР П1.3В1.3 Гипроздрав (ЖД район Останкинский, ул. Бочкова, влд. 11А)\\221049161-ОПР П1.3В1.3 Гипроздрав (ЖД район Останкинский, ул. Бочкова, влд. 11А).doc',  # 29
        '211006810в-ОПР П1(К) НИИПРОЕКТ ГБУ МО (Поликлиника Хотьково, КДЦ).doc',  # 30
        f'{for_automata}Каналка\\ПВ9 - Регина.docx',  # 31
        'ПВ9 исходный.docx',  # 32
        f'{for_automata}Вентиляторы\\В1 от 09.01.23 (ВРАН6-090-Т80).docx',  # 33
        'МО2_ВИР800-045(1)_11(2F)_RD0_9100м3_2400Па копия.docx',  # 34
        f'{for_automata}Каналка\\П5 - Николай.docx',  # 35
        f'{for_automata}Вентиляторы\\МО2_ВИР800-045(1)_11(2F)_RD0_9100м3_2400Па.pdf',  # 36
        f'{for_automata}Веросы\\231028713-ОПР ПВ1 МКС (Производственное здание г.о. Подольск, с. Сынково, 81).doc',  # 37
        f'{for_automata}Веросы\\БЗ Регион\\211022081-ПРМ.doc',  # 38
        f'{for_automata}Вперемешку\\2\\В1 БКТ1.docx',  # 39
        f'{for_automata}Вентиляторы\\ПД1 версия 2.docx',  # 40
        f'{for_automata}Веросы\\231010994-ОПР ПОДЗЕМПРОЕКТ(НПП Исток им. Шокина)П30.2.docx',  # 41
        f'{for_automata}Каналка\\ПВС-1-20 - 20шт..docx',  # 42

        f'{for_automata}ИТП\\223100972б-КОМ.docx',  # 43
        f'{for_automata}ИТП\\213101089в-ННВ.docx',  # 44
        f'{for_automata}ИТП\\233100050б-КОМ.docx',  # 45
        f'{for_automata}ИТП\\233100095а-КОМ.docx',  # 46
        f'{for_automata}ИТП\\233100308-КОМ.docx ',  # 47
        f'{for_automata}ИТП\\233100330-КОМ.docx',  # 48
        f'{for_automata}ИТП\\233100331-КОМ.docx ',  # 49
        f'{for_automata}ИТП\\233100332-КОМ.docx',  # 50
        f'{for_automata}ИТП\\233100406-КОМ.docx',  # 51
        f'{for_automata}ИТП\\233100411-ННВ.docx ',  # 52
        f'{for_automata}ИТП\\233100427-КОМ.docx ',  # 53
        f'{for_automata}ИТП\\233100434-КОМ.docx',  # 54
        f'{for_automata}ИТП\\233100439-КОМ.docx ',  # 55
        f'{for_automata}ИТП\\233100443-КОМ.docx',  # 56
        f'{for_automata}ИТП\\233100447-КОМ.docx ',  # 57
        f'{for_automata}ИТП\\233100474-КОМ.docx',  # 58
        f'{for_automata}ИТП\\233100476-КОМ.pdf',  # 59
        f'{for_automata}ИТП\\233100499-КОМ.pdf',  # 60

        'c:\\Users\\ovchinnikov\\Documents\\Python\\21_AUTOMATA\\_ТЗ_для программы\\Каналка\\В1ап.docx',  # 61
    ]

    print(find_all_systems('В1ап'))
    pass

    time_start = time.perf_counter()
    baba = Blank(all_test_files[61])
    PrintDebugModeOn.check_debug_mode()
    PrintDebugModeOn(baba)
    PrintDebugModeOn(repr(baba))
    PrintDebugModeOn(baba.blank_type)
    PrintDebugModeOn(baba.main_information)
    PrintDebugModeOn(baba.all_avaiable_information)
    PrintDebugModeOn(*baba.ALL_MAIN_INFO, sep='\n')
    PrintDebugModeOn(time.perf_counter() - time_start)
    pass

    PrintDebugModeOn.debug_mode_tumbler()
    PrintDebugModeOn.check_debug_mode()

    time_start = time.perf_counter()
    PrintDebugModeOn(baba)
    PrintDebugModeOn(repr(baba))
    PrintDebugModeOn(baba.blank_type)
    PrintDebugModeOn(baba.main_information)
    PrintDebugModeOn(baba.all_avaiable_information)
    PrintDebugModeOn(*baba.ALL_MAIN_INFO, sep='\n')
    PrintDebugModeOn(time.perf_counter() - time_start)
    pass

    PrintDebugModeOn.debug_mode_tumbler()
    PrintDebugModeOn.check_debug_mode()

    time_start = time.perf_counter()
    PrintDebugModeOn(baba)
    PrintDebugModeOn(repr(baba))
    PrintDebugModeOn(baba.blank_type)
    PrintDebugModeOn(baba.main_information)
    PrintDebugModeOn(baba.all_avaiable_information)
    PrintDebugModeOn(*baba.ALL_MAIN_INFO, sep='\n')
    PrintDebugModeOn(time.perf_counter() - time_start)
    pass

    for test_file in all_test_files[43:]:
    # test_file = all_test_files[43]
        print(test_file)
        dodoco = Document(test_file)
        test_blank = Blank(dodoco)
        print(test_blank, test_blank.all_main_info_text, sep='\n')
        print(test_blank.blank_type)
        print(test_blank.ordered_content)

        print(find_all_systems(test_blank.main_information['Название']))
else:
    PrintDebugModeOn.debug_mode_tumbler()